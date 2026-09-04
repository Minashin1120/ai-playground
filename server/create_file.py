# ============================================================================
# Model file-creation tool ("create_file")
# ----------------------------------------------------------------------------
# Lets chat models create text / markdown / code / PDF / DOCX / XLSX files and
# save them straight into the user's file library.  Generation happens entirely
# server-side (no bubblewrap sandbox is involved), so nothing is left behind in
# the sandbox and code files are written, never executed.
# ============================================================================

# Allowed output extensions for the create_file tool.
_CREATE_FILE_TEXT_EXTS = {
    ".txt", ".text", ".md", ".markdown", ".mdown", ".mkd", ".log", ".rst",
    ".csv", ".tsv", ".json", ".jsonl", ".ndjson", ".xml", ".yaml", ".yml",
    ".toml", ".ini", ".cfg", ".conf", ".env", ".properties",
}
_CREATE_FILE_CODE_EXTS = {
    ".py", ".pyw", ".js", ".mjs", ".cjs", ".ts", ".jsx", ".tsx", ".html",
    ".htm", ".css", ".scss", ".sass", ".less", ".sh", ".bash", ".zsh", ".fish",
    ".sql", ".java", ".kt", ".kts", ".c", ".h", ".cpp", ".hpp", ".cc", ".cs",
    ".go", ".rs", ".rb", ".php", ".swift", ".scala", ".dart", ".lua", ".pl",
    ".r", ".m", ".mm", ".ipynb", ".vue", ".svelte", ".astro", ".dockerfile",
    "dockerfile", ".makefile", "makefile",
}
_CREATE_FILE_DOC_EXTS = {".pdf", ".docx", ".xlsx", ".xlsm"}
_CREATE_FILE_MAX_BYTES = 8 * 1024 * 1024  # max generated file size
_CREATE_FILE_MAX_CONTENT_CHARS = 4 * 1024 * 1024  # max tool content input

_CREATE_FILE_FORMAT_BY_EXT = {
    ".pdf": "pdf", ".docx": "docx", ".xlsx": "xlsx", ".xlsm": "xlsx",
    ".md": "markdown", ".markdown": "markdown",
    ".csv": "csv", ".tsv": "tsv", ".json": "json",
}

# Default extension used when the model supplies an explicit ``format`` but the
# filename has no extension (e.g. create_file(filename="企画書", format="docx")).
_CREATE_FILE_EXT_BY_FORMAT = {
    "text": ".txt",
    "markdown": ".md",
    "code": ".txt",
    "pdf": ".pdf",
    "docx": ".docx",
    "xlsx": ".xlsx",
    "csv": ".csv",
    "tsv": ".tsv",
    "json": ".json",
}

_CREATE_FILE_FORMATS = ("text", "markdown", "code", "pdf", "docx", "xlsx", "csv", "tsv", "json")

# Extensions a user may upload from their own device.  This intentionally
# covers every extension the create_file tool can produce (text / markdown /
# code / PDF / Word / Excel) so that files the model creates on the server can
# be downloaded to the device and uploaded again later (e.g. to attach them in
# another chat) without being rejected.
_UPLOAD_ALLOWED_EXTENSIONS = frozenset({
    '.txt', '.pdf', '.docx', '.png', '.jpg', '.jpeg', '.gif', '.webp',
    '.wav', '.mp3', '.m4a', '.ogg', '.flac', '.webm', '.mp4', '.mov', '.mkv',
    '.avi', '.m4v',
}) | frozenset(_CREATE_FILE_TEXT_EXTS) | frozenset(_CREATE_FILE_CODE_EXTS) | frozenset(_CREATE_FILE_DOC_EXTS)

# Extensions whose contents are text (or code) and can be read as plain text
# by the chat models when attached.  Used to detect "text-like" attachments
# even when the platform MIME map reports a non-text MIME (e.g. .json, .yaml).
_TEXT_LIKE_UPLOAD_EXTS = frozenset(_CREATE_FILE_TEXT_EXTS) | frozenset(_CREATE_FILE_CODE_EXTS)


def _create_file_allowed_ext(ext):
    return ext in _CREATE_FILE_TEXT_EXTS or ext in _CREATE_FILE_CODE_EXTS or ext in _CREATE_FILE_DOC_EXTS


def _infer_create_file_format(filename):
    ext = os.path.splitext(str(filename or ""))[1].lower()
    if ext in _CREATE_FILE_FORMAT_BY_EXT:
        return _CREATE_FILE_FORMAT_BY_EXT[ext]
    if ext in _CREATE_FILE_DOC_EXTS:
        return "pdf" if ext == ".pdf" else ("docx" if ext == ".docx" else "xlsx")
    if ext in _CREATE_FILE_CODE_EXTS:
        return "code"
    return "text"


def _sanitize_create_file_base_name(filename):
    """Return a safe ASCII file base-name that preserves the original extension.

    ``secure_filename()`` strips non-ASCII characters, so Japanese names such
    as ``企画書.docx`` become ``docx`` (the leading dot is also removed),
    which makes the extension check fail with an empty extension.  The
    extension is therefore read from the original basename and re-appended to
    the sanitized stem (same approach as ``_sanitize_upload_filename``).  The
    returned name is used for extension detection and on-disk storage;
    ``_execute_create_file_tool`` keeps the original name for user display.
    """
    raw = str(filename or "").strip()
    raw = raw.replace("\\", "/")
    base = raw.rsplit("/", 1)[-1].strip()
    ext = os.path.splitext(base)[1].lower()
    stem = secure_filename(base) or ""
    if ext and not stem.endswith(ext):
        stem = secure_filename(os.path.splitext(base)[0]) or "file"
        return f"{stem}{ext}"
    return stem or "file"


def _inline_library_images_for_create_file(markdown_text, user_id):
    """Rewrite markdown image references so embedded documents can include
    images from the user's library (/files/...) or data URIs.

    Returns (markdown_text, warnings).
    """
    if not markdown_text:
        return markdown_text, []
    warnings = []
    # Matches ![alt](src "title") and ![alt](src)
    pattern = re.compile(r"!\[([^\]]*)\]\(([^)\s]+)(?:\s+[\"'].*?[\"'])?\)")
    consumed_srcs = []

    def _replace(match):
        alt = match.group(1)
        src = match.group(2).strip()
        if src.startswith("data:"):
            return match.group(0)
        resolved = None
        # Support /files/<user_id>/<name> and files/<user_id>/<name> and relative library refs.
        norm_candidates = []
        src_clean = src.split("?")[0].split("#")[0]
        if src_clean.startswith("/files/"):
            norm_candidates.append(src_clean[len("/files/"):])
        elif src_clean.startswith("files/"):
            norm_candidates.append(src_clean[len("files/"):])
        else:
            norm_candidates.append(src_clean)
        for cand in norm_candidates:
            norm = _normalize_upload_ref(cand)
            if not norm:
                continue
            info = _get_file_disk_info(norm)
            if not info or not info.get("exists"):
                continue
            data = _load_user_file_bytes(norm, info)
            if not data:
                continue
            mime = _normalize_media_mime(norm, mimetypes.guess_type(norm)[0] or "application/octet-stream")
            b64 = base64.b64encode(data).decode("ascii")
            resolved = f"data:{mime};base64,{b64}"
            break
        if resolved is None:
            warnings.append(f"画像参照を解決できませんでした: {src[:120]}")
            return match.group(0)
        consumed_srcs.append(src)
        return f"![{alt}]({resolved})"

    new_text = pattern.sub(_replace, markdown_text)
    return new_text, warnings


def _build_created_pdf_bytes(content_md, title):
    """Render markdown content to a multi-page PDF using WeasyPrint."""
    import markdown as _md

    html_body = _md.markdown(
        content_md or "",
        extensions=["tables", "fenced_code", "sane_lists", "nl2br"],
    )
    safe_title = html.escape(str(title or "Document").strip() or "Document")
    document_html = f"""<!doctype html>
<html lang="ja">
<head>
<meta charset="utf-8">
<title>{safe_title}</title>
<style>
@page {{
  size: A4;
  margin: 16mm 16mm 18mm;
  @bottom-right {{
    content: "Page " counter(page) " / " counter(pages);
    color: #6b7280;
    font-size: 8.5pt;
  }}
}}
html {{ color-scheme: light; background: #ffffff; }}
body {{
  margin: 0;
  background: #ffffff;
  color: #1f2937;
  font-family: "IPAPGothic", "IPAGothic", "Droid Sans Fallback", sans-serif;
  font-size: 10.5pt;
  line-height: 1.55;
  overflow-wrap: anywhere;
  word-break: normal;
}}
h1 {{ font-size: 18pt; margin: 0 0 4mm; padding-bottom: 2mm; border-bottom: 1.5pt solid #e5e7eb; break-after: avoid; }}
h2 {{ font-size: 15pt; margin: 5mm 0 2.5mm; break-after: avoid; }}
h3 {{ font-size: 13pt; margin: 4mm 0 2mm; break-after: avoid; }}
h4, h5, h6 {{ font-size: 11pt; margin: 3mm 0 1.5mm; break-after: avoid; }}
p {{ margin: 0 0 0.85em; }}
ul, ol {{ margin: 0 0 0.85em; padding-left: 6mm; }}
li {{ margin: 0 0 0.2em; }}
blockquote {{ margin: 0 0 0.85em; padding: 2mm 4mm; border-left: 2pt solid #9ca3af; color: #4b5563; }}
img {{ display: block; max-width: 100%; height: auto; margin: 0.9em auto; }}
pre {{
  background: #f3f4f6;
  padding: 2.5mm 3mm;
  border-radius: 2mm;
  font-family: "DejaVu Sans Mono", monospace;
  font-size: 8.5pt;
  white-space: pre-wrap;
  overflow-wrap: anywhere;
}}
code {{ font-family: "DejaVu Sans Mono", monospace; font-size: 8.8pt; }}
pre code {{ font-size: 8.5pt; background: transparent; padding: 0; }}
table {{
  max-width: 100%;
  margin: 1em 0;
  border-collapse: collapse;
  font-size: 9pt;
}}
thead {{ display: table-header-group; }}
tr {{ break-inside: avoid; }}
th, td {{ border: 0.6pt solid #d1d5db; padding: 2.2mm 2.5mm; text-align: left; }}
th {{ background: #f3f4f6; }}
hr {{ margin: 1em 0; border: 0; border-top: 0.7pt solid #d1d5db; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""
    weasyprint_bin = os.path.join(os.path.dirname(sys.executable), "weasyprint")
    if not os.path.isfile(weasyprint_bin):
        raise RuntimeError("weasyprint_command_not_found")
    completed = subprocess.run(
        [
            weasyprint_bin,
            "--quiet",
            "--presentational-hints",
            "--optimize-images",
            "--jpeg-quality", "92",
            "--dpi", "180",
            "--timeout", "5",
            "--allowed-protocols", "data",
            "--no-http-redirects",
            "-", "-",
        ],
        input=document_html.encode("utf-8"),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=90,
        check=False,
    )
    if completed.returncode != 0:
        stderr = completed.stderr.decode("utf-8", errors="replace").strip()
        raise RuntimeError(f"weasyprint_failed: {stderr[-1000:]}")
    pdf_bytes = completed.stdout
    if not isinstance(pdf_bytes, bytes) or not pdf_bytes.startswith(b"%PDF"):
        raise RuntimeError("weasyprint_invalid_pdf")
    return pdf_bytes


def _build_created_docx_bytes(content_md, title):
    """Render markdown content to a .docx file using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(str(title or "Document").strip() or "Document", level=0)

    lines = (content_md or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    n = len(lines)

    def _add_runs(paragraph, text):
        import re as _re
        token_re = _re.compile(r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|`[^`]+`)")
        for token in token_re.split(text):
            if not token:
                continue
            if token.startswith("**") and token.endswith("**") and len(token) > 4:
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("__") and token.endswith("__") and len(token) > 4:
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*") and token.endswith("*") and len(token) > 2:
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            elif token.startswith("`") and token.endswith("`") and len(token) > 2:
                run = paragraph.add_run(token[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            else:
                paragraph.add_run(token)

    def _parse_table_rows(header_line, body_lines):
        def _split_row(line):
            line = line.strip()
            if line.startswith("|"):
                line = line[1:]
            if line.endswith("|"):
                line = line[:-1]
            return [cell.strip() for cell in line.split("|")]

        rows = [_split_row(header_line)]
        for bl in body_lines:
            bl = bl.strip()
            if not bl:
                continue
            if set(bl.replace("|", "").replace(" ", "").replace(":", "")) == set("-"):
                continue
            rows.append(_split_row(bl))
        return rows

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```") or stripped.startswith("~~~"):
            fence_char = stripped[0]
            i += 1
            code_lines = []
            while i < n:
                if lines[i].strip().startswith(fence_char * 3):
                    i += 1
                    break
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.2)
            continue

        # Table block
        if stripped.startswith("|") and i + 1 < n:
            body_lines = []
            j = i + 1
            while j < n:
                bj = lines[j].strip()
                if bj.startswith("|") or (j == i + 1 and set(bj.replace("|", "").replace(" ", "").replace(":", "")) == set("-")):
                    body_lines.append(bj)
                    j += 1
                else:
                    break
            if body_lines:
                rows = _parse_table_rows(stripped, body_lines)
                if len(rows) > 1:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = "Table Grid"
                    for r_idx, row in enumerate(rows):
                        for c_idx in range(min(len(row), len(rows[0]))):
                            cell = table.cell(r_idx, c_idx)
                            cell.text = ""
                            _add_runs(cell.paragraphs[0], row[c_idx])
                    doc.add_paragraph()
                    i = j
                    continue

        # Headings
        h_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if h_match and len(h_match.group(1)) <= 6:
            level = len(h_match.group(1))
            doc.add_heading(h_match.group(2).strip(), level=level)
            i += 1
            continue

        # Horizontal rule
        if re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", stripped):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # Unordered list
        if re.match(r"^[-*+]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^[-*+]\s+", "", stripped))
            i += 1
            continue

        # Ordered list
        if re.match(r"^\d+[.)]\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\d+[.)]\s+", "", stripped))
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            _add_runs(p, stripped.lstrip("> ").strip())
            p.paragraph_format.left_indent = Inches(0.2)
            i += 1
            continue

        # Image
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)\s]+)\)\s*$", stripped)
        if img_match:
            src = img_match.group(2).strip()
            if src.startswith("data:"):
                try:
                    header, _, b64 = src.partition(",")
                    img_bytes = base64.b64decode(b64)
                    from io import BytesIO as _BytesIO
                    doc.add_picture(_BytesIO(img_bytes), width=Inches(5.5))
                except Exception:
                    p = doc.add_paragraph(img_match.group(1) or "画像")
            else:
                doc.add_paragraph(img_match.group(1) or src)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_runs(p, stripped)
        i += 1

    from io import BytesIO as _BytesIO2
    buf = _BytesIO2()
    doc.save(buf)
    return buf.getvalue()


# Marker used both by _extract_xlsx_as_tsv and _build_created_xlsx_bytes so a
# workbook with multiple sheets can round-trip through the TSV text form.
_XLSX_SHEET_MARKER = "# Sheet: "


def _escape_tsv_cell(value):
    """Escape a spreadsheet cell value for the TSV text form.

    Tabs, newlines and backslashes inside a cell are escaped so the text stays
    one logical row per spreadsheet row and can round-trip losslessly.
    """
    return (
        str(value)
        .replace("\\", "\\\\")
        .replace("\t", "\\t")
        .replace("\r", "\\r")
        .replace("\n", "\\n")
    )


def _unescape_tsv_cell(value):
    """Undo _escape_tsv_cell when rebuilding a workbook from TSV text."""
    out = []
    s = str(value)
    i = 0
    while i < len(s):
        if s[i] == "\\" and i + 1 < len(s):
            nxt = s[i + 1]
            if nxt == "n":
                out.append("\n")
            elif nxt == "t":
                out.append("\t")
            elif nxt == "r":
                out.append("\r")
            elif nxt == "\\":
                out.append("\\")
            else:
                out.append(s[i])
            i += 2
        else:
            out.append(s[i])
            i += 1
    return "".join(out)


def _build_created_xlsx_bytes(content_tsv):
    """Render TSV content to an .xlsx workbook using openpyxl.

    Lines starting with ``# Sheet: <name>`` start a new worksheet, so workbooks
    produced from _extract_xlsx_as_tsv (which emits that marker for every sheet
    after the first) can be rebuilt with their sheet structure preserved.
    Cells may use the \\n / \\t / \\\\ escapes produced by _escape_tsv_cell.
    """
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    lines = (content_tsv or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    row_index = 1
    for line in lines:
        if line.startswith(_XLSX_SHEET_MARKER):
            sheet_name = line[len(_XLSX_SHEET_MARKER):].strip()[:31] or "Sheet"
            ws = wb.create_sheet(title=sheet_name)
            row_index = 1
            continue
        if not line.strip():
            continue
        cells = line.split("\t")
        for col_index, cell in enumerate(cells, start=1):
            ws.cell(row=row_index, column=col_index, value=_unescape_tsv_cell(cell))
        row_index += 1
    from io import BytesIO as _BytesIO3
    buf = _BytesIO3()
    wb.save(buf)
    return buf.getvalue()


def _excel_column_letters(count):
    """Return the first ``count`` Excel column letters (A, B, ..., Z, AA, AB, ...)."""
    letters = []
    for i in range(1, int(count) + 1):
        n = i
        label = ""
        while n > 0:
            n, rem = divmod(n - 1, 26)
            label = chr(65 + rem) + label
        letters.append(label)
    return letters


def _extract_xlsx_as_tsv(data, max_sheets=8, max_cells=500000, include_column_headers=False):
    """Read .xlsx/.xlsm bytes and return a tab-separated text representation.

    Every worksheet is rendered as one TSV row per spreadsheet row.  When the
    workbook has more than one sheet, a ``# Sheet: <name>`` marker line is
    emitted before each sheet.  When ``include_column_headers`` is set, a row of
    Excel column letters (A, B, C, ...) is emitted at the top of each sheet so a
    model can derive cell addresses (e.g. AD15) to feed back into the edit_file
    tool.  Cells that contain tabs / newlines / backslashes are escaped with
    _escape_tsv_cell so the text stays one row per spreadsheet row.
    Returns None when the bytes cannot be parsed.
    """
    if not data:
        return None
    try:
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
    except Exception:
        return None
    out_lines = []
    cell_count = 0
    sheets = list(wb.worksheets)[:max_sheets]
    for idx, ws in enumerate(sheets):
        # The first sheet maps onto the default "Sheet1" created by
        # _build_created_xlsx_bytes, so only later sheets need a marker.
        if idx > 0 and len(sheets) > 1:
            out_lines.append(f"{_XLSX_SHEET_MARKER}{ws.title}")
        try:
            row_iter = ws.iter_rows(values_only=True)
        except Exception:
            continue
        sheet_lines = []
        max_cols = 0
        for row in row_iter:
            if row is None:
                continue
            cells = []
            for value in row:
                if value is None:
                    cells.append("")
                elif isinstance(value, datetime):
                    if value.hour or value.minute or value.second:
                        cells.append(value.strftime("%Y-%m-%d %H:%M:%S"))
                    else:
                        cells.append(value.strftime("%Y-%m-%d"))
                elif isinstance(value, (int, float)) and not isinstance(value, bool):
                    # Keep numbers compact (avoid 1.0 for integer cells).
                    if isinstance(value, float) and value.is_integer():
                        cells.append(str(int(value)))
                    else:
                        cells.append(str(value))
                else:
                    cells.append(_escape_tsv_cell(value))
                cell_count += 1
                if cell_count > max_cells:
                    break
            while cells and cells[-1] == "":
                cells.pop()
            max_cols = max(max_cols, len(cells))
            sheet_lines.append("\t".join(cells))
            if cell_count > max_cells:
                break
        if include_column_headers and max_cols:
            out_lines.append("\t".join(_excel_column_letters(max_cols)))
        out_lines.extend(sheet_lines)
        if cell_count > max_cells:
            break
    return "\n".join(out_lines)


def _create_file_generated_bytes(filename, format_name, content, user_id):
    """Return the generated file bytes for the given create_file arguments."""
    if format_name in ("text", "csv", "tsv", "json"):
        return str(content or "").encode("utf-8")
    if format_name in ("markdown", "code"):
        return str(content or "").encode("utf-8")
    if format_name == "pdf":
        resolved_md, _warnings = _inline_library_images_for_create_file(str(content or ""), user_id)
        return _build_created_pdf_bytes(resolved_md, os.path.splitext(filename)[0])
    if format_name == "docx":
        resolved_md, _warnings = _inline_library_images_for_create_file(str(content or ""), user_id)
        return _build_created_docx_bytes(resolved_md, os.path.splitext(filename)[0])
    if format_name == "xlsx":
        return _build_created_xlsx_bytes(str(content or ""))
    raise ValueError(f"Unsupported create_file format: {format_name}")


def _execute_create_file_tool(user_id, args, encrypt):
    """Execute the create_file tool server-side.  Returns a dict result.

    result = {"ok": True, "filename": ..., "url": ...} or
             {"ok": False, "error": ...}
    """
    if not isinstance(args, dict):
        return {"ok": False, "error": "create_fileの引数が不正です。"}
    filename_raw = str(args.get("filename") or "").strip()
    if not filename_raw:
        return {"ok": False, "error": "filename は必須です。"}
    safe_base = _sanitize_create_file_base_name(filename_raw)
    ext = os.path.splitext(safe_base)[1].lower()
    explicit_format = str(args.get("format") or "").strip().lower()
    if explicit_format not in _CREATE_FILE_FORMATS:
        explicit_format = ""
    format_name = explicit_format or _infer_create_file_format(safe_base)
    if not _create_file_allowed_ext(ext):
        # An explicit format lets the model create a file even when the
        # filename has no (recognized) extension; the format's default
        # extension is used for the on-disk name.
        if explicit_format and explicit_format in _CREATE_FILE_EXT_BY_FORMAT and not ext:
            ext = _CREATE_FILE_EXT_BY_FORMAT[explicit_format]
            safe_base = f"{os.path.splitext(safe_base)[0]}{ext}"
        else:
            return {"ok": False, "error": f"対応していないファイル形式です: {ext or '(拡張子なし)'}"}
    # User-facing name keeps the original (possibly non-ASCII) basename while
    # the extension is normalized to the detected lowercase ext, so the model
    # can present/download the file under its original Japanese name.
    display_name = _sanitize_file_display_name(filename_raw) or "file"
    display_ext = os.path.splitext(display_name)[1]
    if not display_ext:
        display_name = f"{display_name}{ext}"
    elif display_ext.lower() != ext:
        display_name = f"{os.path.splitext(display_name)[0]}{ext}"
    content = args.get("content")
    if content is None:
        return {"ok": False, "error": "content は必須です。"}
    content = str(content)
    if len(content) > _CREATE_FILE_MAX_CONTENT_CHARS:
        return {"ok": False, "error": "ファイル内容が大きすぎます。"}
    try:
        data = _create_file_generated_bytes(display_name, format_name, content, user_id)
    except Exception as exc:
        logger.warning("create_file generation failed: %s", exc)
        return {"ok": False, "error": f"ファイル生成に失敗しました: {str(exc)[:200]}"}
    if not isinstance(data, (bytes, bytearray)) or not data:
        return {"ok": False, "error": "生成されたファイルが空です。"}
    if len(data) > _CREATE_FILE_MAX_BYTES:
        return {"ok": False, "error": "生成されたファイルが大きすぎます。"}

    def _make_unique_filename():
        return f"{os.path.splitext(safe_base)[0]}_{int(time.time())}_{os.urandom(3).hex()}{ext}"

    try:
        fname, url = _save_user_generated_bytes_verified(user_id, bytes(data), _make_unique_filename, bool(encrypt))
    except StorageLimitError as exc:
        return {"ok": False, "error": str(exc)}
    except Exception as exc:
        logger.warning("create_file save failed: %s", exc)
        return {"ok": False, "error": f"ライブラリへの保存に失敗しました: {str(exc)[:200]}"}
    return {"ok": True, "filename": fname, "display_name": display_name, "url": url, "size": len(data)}


def _create_file_tool_result_text(result, tool_name="create_file", action="作成"):
    """Convert a file-tool result dict into the text fed back to the model.

    Shared by create_file (default) and edit_file (pass tool_name="edit_file",
    action="編集").  A ``note`` key on the result (e.g. partial-edit warnings)
    is appended when present.
    """
    if result.get("ok"):
        text = (
            f"ファイルを{action}しました。\n"
            f"filename: {result.get('display_name')}\n"
            f"url: {result.get('url')}\n"
            f"ユーザーにはこのURLでファイルをダウンロードできるリンクを提供してください。"
        )
        extra_files = result.get("extra_files")
        if extra_files:
            for extra in extra_files:
                text += (
                    f"\n追加ファイル（{action}したファイルから変換・生成したもの）: "
                    f"{extra.get('display_name')} - {extra.get('url')}"
                )
        note = result.get("note")
        if note:
            text = f"{text}\n{note}"
        return text
    return f"{tool_name} エラー: {result.get('error')}"


def _build_create_file_tool_schema():
    """Return the OpenAI-compatible tool schema for create_file."""
    return {
        "type": "function",
        "function": {
            "name": "create_file",
            "description": (
                "テキスト・Markdown・コード・PDF・Word (docx)・Excel (xlsx) などのファイルを作成し、"
                "ユーザーのファイルライブラリに保存します。保存後のURLが返るので、それを回答中でリンクとして提示してください。"
                "PDF/DOCX の content は Markdown 形式（見出し・段落・表・画像）。"
                "XLSX の content は TSV（タブ区切り、1行目がヘッダー）。"
                "コードファイルは実行せず、そのままテキストとして保存されます。"
                "添付されている既存ファイルの編集を求められた場合は、新規作成せず edit_file ツールを使用してください。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "作成するファイル名（拡張子含む）。例: report.txt, memo.md, script.py, document.pdf, data.xlsx",
                    },
                    "format": {
                        "type": "string",
                        "enum": list(_CREATE_FILE_FORMATS),
                        "description": "ファイル形式。省略時は拡張子から自動判定されます。",
                    },
                    "content": {
                        "type": "string",
                        "description": (
                            "ファイルの内容。text/markdown/code はそのままのテキスト（改行を含めて保存）。"
                            "pdf/docx は Markdown 形式。xlsx は TSV。"
                        ),
                    },
                },
                "required": ["filename", "content"],
                "additionalProperties": False,
            },
        },
    }


def _resolve_attached_file_ref(source, loaded_files=None, history=None, user_id=None, thread_id=None):
    """Resolve a model-provided ``source`` reference to a target file.

    Matches against the user-facing send name, the stored rel path, and their
    basenames (case-insensitive) across:
    1. The current turn's loaded attachments (``loaded_files``).
    2. Files attached or generated across messages in the conversation ``history``
       (both user uploads and assistant-generated files).
       If an original file has been edited into newer versions in the thread,
       prioritizes the latest edited version so follow-up edits are cumulative.
    3. User's file library / FileCache / uploads disk for ``user_id``.

    Returns the matching entry dictionary or ``None``.
    """
    if not source:
        return None
    src = str(source).strip()
    if not src:
        return None

    # Strip URL formatting like /files/... or markdown link syntax [name](/files/path)
    url_m = re.search(r'/files/([^)\s?#]+)', src)
    if url_m:
        src = url_m.group(1).strip()
    elif src.startswith('[') and '](' in src:
        inner_m = re.search(r'\((?:/files/)?([^)\s?#]+)\)', src)
        if inner_m:
            src = inner_m.group(1).strip()

    # Strip leading [File: ...] wrapper if present
    file_tag_m = re.match(r'^\[(?:File|file):\s*(.*?)\s*\]$', src)
    if file_tag_m:
        src = file_tag_m.group(1).strip()

    src_lower = src.lower()
    src_base = os.path.basename(src).lower()
    src_stem = os.path.splitext(src_base)[0].lower()
    src_ext = os.path.splitext(src_base)[1].lower()

    # 1. Search loaded_files (current message attachments)
    if loaded_files:
        for fi in loaded_files:
            candidates = [
                str(fi.get('send_name') or ''),
                str(fi.get('name') or ''),
                str(fi.get('path') or ''),
            ]
            for name in candidates:
                if not name:
                    continue
                name_lower = name.lower()
                name_base = os.path.basename(name).lower()
                if (
                    name_lower == src_lower
                    or name_base == src_lower
                    or name_lower == src_base
                    or name_base == src_base
                ):
                    return fi

    # 2. Search conversation history (from newest message to oldest)
    if history:
        history_candidates = []
        seen_paths = set()
        for m in reversed(history):
            role = str(m.get('role') or '').lower()
            # 2a. image_url (which holds JSON list or string of uploaded / generated file refs)
            raw_urls = m.get('image_url')
            if raw_urls:
                try:
                    ref_list = json.loads(raw_urls)
                except Exception:
                    ref_list = raw_urls
                if not isinstance(ref_list, list):
                    ref_list = [ref_list]
                for ref in ref_list:
                    norm = _normalize_upload_ref(ref)
                    if norm and norm not in seen_paths:
                        if user_id is not None and not norm.startswith(f"{user_id}/"):
                            continue
                        info = _get_file_disk_info(norm)
                        if info.get('exists'):
                            seen_paths.add(norm)
                            history_candidates.append({
                                'path': norm,
                                'name': norm,
                                'send_name': os.path.basename(norm),
                                'is_assistant': (role == 'assistant'),
                                'mtime': info.get('mtime', 0),
                            })
            # 2b. markdown links in message content: e.g. [display_name](/files/user_id/filename)
            content = m.get('content') or ''
            if content:
                for link_name, link_url in re.findall(r'\[([^\]]+)\]\((?:/files/)?([^\)\s?#]+)\)', content):
                    norm = _normalize_upload_ref(link_url)
                    if not norm:
                        continue
                    if user_id is not None and not norm.startswith(f"{user_id}/"):
                        continue
                    # Update display name on existing candidate if found
                    updated = False
                    for cand in history_candidates:
                        if cand['path'] == norm:
                            if link_name.strip():
                                cand['send_name'] = link_name.strip()
                            updated = True
                            break
                    if not updated and norm not in seen_paths:
                        info = _get_file_disk_info(norm)
                        if info.get('exists'):
                            seen_paths.add(norm)
                            history_candidates.append({
                                'path': norm,
                                'name': norm,
                                'send_name': link_name.strip() or os.path.basename(norm),
                                'is_assistant': (role == 'assistant'),
                                'mtime': info.get('mtime', 0),
                            })

        # 2c. Exact matches in history candidates (newest first)
        for cand in history_candidates:
            candidates = [
                str(cand.get('send_name') or ''),
                str(cand.get('name') or ''),
                str(cand.get('path') or ''),
            ]
            for name in candidates:
                if not name:
                    continue
                name_lower = name.lower()
                name_base = os.path.basename(name).lower()
                if (
                    name_lower == src_lower
                    or name_base == src_lower
                    or name_lower == src_base
                    or name_base == src_base
                ):
                    return cand

        # 2d. Stem / prefix match (e.g. source is the original filename or display name,
        # but a newer assistant-edited version exists in history)
        for cand in history_candidates:
            cand_base = os.path.basename(cand.get('path') or '').lower()
            cand_stem = os.path.splitext(cand_base)[0]
            cand_ext = os.path.splitext(cand_base)[1].lower()
            cand_send_base = os.path.basename(cand.get('send_name') or '').lower()
            cand_send_stem = os.path.splitext(cand_send_base)[0]

            if src_ext and cand_ext and src_ext != cand_ext:
                continue

            if (
                cand_stem.startswith(src_stem)
                or src_stem.startswith(cand_stem)
                or cand_send_stem == src_stem
                or cand_send_base == src_base
            ):
                return cand

    # 3. Fallback: Search user's library / disk files for user_id
    if user_id is not None:
        try:
            cache_rows = FileCache.query.filter_by(user_id=user_id).all()
            for row in cache_rows:
                cand_path = row.path
                if not cand_path:
                    continue
                cand_base = os.path.basename(cand_path).lower()
                cand_stem = os.path.splitext(cand_base)[0]
                if (
                    cand_path.lower() == src_lower
                    or cand_base == src_lower
                    or cand_base == src_base
                    or (src_stem and cand_stem == src_stem)
                ):
                    norm = _normalize_upload_ref(cand_path)
                    if norm:
                        info = _get_file_disk_info(norm)
                        if info.get('exists'):
                            return {
                                'path': norm,
                                'name': norm,
                                'send_name': os.path.basename(norm),
                                'mtime': info.get('mtime', 0),
                            }
        except Exception:
            pass

        try:
            user_dir = os.path.join(app.config.get('UPLOAD_FOLDER', 'instance/uploads'), str(user_id))
            if os.path.isdir(user_dir):
                for fname in os.listdir(user_dir):
                    fname_clean = fname[:-4] if fname.endswith('.enc') else fname
                    fname_lower = fname_clean.lower()
                    fname_base = os.path.basename(fname_clean).lower()
                    fname_stem = os.path.splitext(fname_base)[0]
                    if (
                        fname_lower == src_lower
                        or fname_base == src_base
                        or (src_stem and fname_stem == src_stem)
                    ):
                        rel_p = f"{user_id}/{fname_clean}"
                        info = _get_file_disk_info(rel_p)
                        if info.get('exists'):
                            return {
                                'path': rel_p,
                                'name': rel_p,
                                'send_name': fname_clean,
                                'mtime': info.get('mtime', 0),
                            }
        except Exception:
            pass

    return None


