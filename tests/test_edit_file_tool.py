import os
from io import BytesIO
from pathlib import Path
import unittest
from unittest.mock import patch

from tests.app_source import read_app_source
os.environ.setdefault("FLASK_SECRET_KEY", "edit-file-tool-test-secret")
os.environ.setdefault("DATABASE_URL", "sqlite:////tmp/ai-chat-edit-file-tool-tests.db")
os.environ.setdefault("REDIS_URL", "redis://127.0.0.1:6399/15")
os.environ.setdefault("RUN_SCHEMA_MIGRATIONS", "0")
os.environ.setdefault("VERBOSE_DEBUG_LOGS", "0")

import app as target


APP_ROOT = Path(__file__).resolve().parents[1]


class XlsxTextExtractionTests(unittest.TestCase):
    def _make_xlsx(self, content_tsv, sheet_name="Sheet1"):
        return target._build_created_xlsx_bytes(content_tsv)

    def test_extract_single_sheet_xlsx_to_tsv(self):
        data = self._make_xlsx("name\tvalue\napple\t3\nbanana\t5\n")
        text = target._extract_xlsx_as_tsv(data)
        self.assertIsNotNone(text)
        self.assertIn("name\tvalue", text)
        self.assertIn("apple\t3", text)
        self.assertIn("banana\t5", text)

    def test_extract_multi_sheet_xlsx_emits_sheet_markers(self):
        data = self._make_xlsx(
            "a\tb\n1\t2\n# Sheet: Sheet2\nx\ty\n3\t4\n"
        )
        text = target._extract_xlsx_as_tsv(data)
        self.assertIsNotNone(text)
        self.assertIn("# Sheet: Sheet2", text)
        self.assertIn("3\t4", text)

    def test_extract_round_trips_through_builder(self):
        data = self._make_xlsx(
            "a\tb\n1\t2\n# Sheet: Sheet2\nx\ty\n3\t4\n"
        )
        text = target._extract_xlsx_as_tsv(data)
        rebuilt = target._build_created_xlsx_bytes(text)
        self.assertTrue(rebuilt.startswith(b"PK"))
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(rebuilt), read_only=True, data_only=True)
        self.assertEqual(wb.sheetnames, ["Sheet1", "Sheet2"])

    def test_extract_invalid_bytes_returns_none(self):
        self.assertIsNone(target._extract_xlsx_as_tsv(b"not an xlsx file"))
        self.assertIsNone(target._extract_xlsx_as_tsv(None))

    def test_extract_with_column_headers(self):
        data = self._make_xlsx("name\tvalue\napple\t3\n")
        text = target._extract_xlsx_as_tsv(data, include_column_headers=True)
        self.assertIsNotNone(text)
        first_line = text.split("\n")[0]
        self.assertEqual(first_line, "A\tB")
        self.assertIn("apple\t3", text)

    def test_excel_column_letters(self):
        self.assertEqual(target._excel_column_letters(3), ["A", "B", "C"])
        self.assertEqual(target._excel_column_letters(28)[-2:], ["AA", "AB"])

    def test_multiline_cells_escape_and_round_trip(self):
        data = self._make_xlsx("a\tb\none\ttwo\\nlines\nthree\ttab\\tsep\n")
        text = target._extract_xlsx_as_tsv(data)
        self.assertIsNotNone(text)
        self.assertIn("two\\nlines", text)
        self.assertIn("tab\\tsep", text)
        rebuilt = target._build_created_xlsx_bytes(text)
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(rebuilt), read_only=True, data_only=True)
        ws = wb.worksheets[0]
        rows = list(ws.iter_rows(values_only=True))
        self.assertEqual(rows[1][1], "two\nlines")
        self.assertEqual(rows[2][1], "tab\tsep")

    def test_build_xlsx_unescapes_cells(self):
        data = target._build_created_xlsx_bytes("a\tb\nx\ty\\\\z\\nq\n")
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
        ws = wb.worksheets[0]
        rows = list(ws.iter_rows(values_only=True))
        self.assertEqual(rows[1][1], "y\\z\nq")

    def test_build_xlsx_handles_sheet_marker(self):
        data = target._build_created_xlsx_bytes(
            "a\tb\n1\t2\n# Sheet: Second\nc\td\n"
        )
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
        self.assertEqual(wb.sheetnames, ["Sheet1", "Second"])


class EditFileToolTests(unittest.TestCase):
    def _make_styled_xlsx(self):
        """Return an xlsx with a filled, non-adjacent cell to verify style retention."""
        from openpyxl import Workbook
        from openpyxl.styles import PatternFill
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["B2"] = "旧値"
        ws["B2"].fill = PatternFill(start_color="FFFFFF00", end_color="FFFFFF00", fill_type="solid")
        ws["C3"] = "保持"
        buf = BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_resolve_matches_send_name_path_and_basename(self):
        loaded = [{
            "send_name": "予定表.xlsx",
            "name": "1/1234_abcd.xlsx",
            "path": "1/1234_abcd.xlsx",
        }]
        self.assertEqual(
            target._resolve_attached_file_ref("予定表.xlsx", loaded), loaded[0]
        )
        self.assertEqual(
            target._resolve_attached_file_ref("1/1234_abcd.xlsx", loaded), loaded[0]
        )
        self.assertEqual(
            target._resolve_attached_file_ref("1234_abcd.xlsx", loaded), loaded[0]
        )
        self.assertEqual(
            target._resolve_attached_file_ref("予定表.XLSX", loaded), loaded[0]
        )

    def test_resolve_returns_none_when_no_match(self):
        loaded = [{"send_name": "予定表.xlsx", "name": "1/1234_abcd.xlsx"}]
        self.assertIsNone(target._resolve_attached_file_ref("missing.csv", loaded))
        self.assertIsNone(target._resolve_attached_file_ref("", loaded))
        self.assertIsNone(target._resolve_attached_file_ref("a.csv", None))

    def test_resolve_from_conversation_history(self):
        history = [
            {"role": "user", "content": "ファイルです", "image_url": '["1/1788181323_fc48c741.xlsx"]'},
            {"role": "assistant", "content": "\n📄 **ファイルを編集しました:** [1788181323_fc48c741.xlsx](/files/1/1788181323_fc48c741_1788190840_b8e900.xlsx)\n", "image_url": '["1/1788181323_fc48c741_1788190840_b8e900.xlsx"]'},
        ]
        with patch("app._get_file_disk_info", return_value={"exists": True, "size": 100}):
            # Follow-up prompt with empty loaded_files referring to original name
            resolved = target._resolve_attached_file_ref(
                "1788181323_fc48c741.xlsx", loaded_files=[], history=history, user_id=1
            )
            self.assertIsNotNone(resolved)
            self.assertEqual(resolved["path"], "1/1788181323_fc48c741_1788190840_b8e900.xlsx")

            # Referring to full URL
            resolved_url = target._resolve_attached_file_ref(
                "/files/1/1788181323_fc48c741_1788190840_b8e900.xlsx", loaded_files=[], history=history, user_id=1
            )
            self.assertIsNotNone(resolved_url)
            self.assertEqual(resolved_url["path"], "1/1788181323_fc48c741_1788190840_b8e900.xlsx")

            # Referring to markdown link
            resolved_md = target._resolve_attached_file_ref(
                "[1788181323_fc48c741.xlsx](/files/1/1788181323_fc48c741_1788190840_b8e900.xlsx)", loaded_files=[], history=history, user_id=1
            )
            self.assertIsNotNone(resolved_md)
            self.assertEqual(resolved_md["path"], "1/1788181323_fc48c741_1788190840_b8e900.xlsx")

    def test_execute_edit_file_in_followup_prompt(self):
        original = self._make_styled_xlsx()
        history = [
            {"role": "user", "content": "編集して", "image_url": '["1/1788181323_fc48c741.xlsx"]'},
            {"role": "assistant", "content": "編集しました [1788181323_fc48c741.xlsx](/files/1/1788181323_fc48c741_1788190840_b8e900.xlsx)", "image_url": '["1/1788181323_fc48c741_1788190840_b8e900.xlsx"]'},
        ]
        # In follow-up prompt, user uploads a screenshot only
        loaded_screenshot = [{"send_name": "screen.png", "name": "1/screen.png", "path": "1/screen.png"}]
        with patch("app._get_file_disk_info", return_value={"exists": True, "size": len(original)}), \
             patch("app._load_user_file_bytes", return_value=original), \
             patch("app._save_user_generated_bytes_verified") as save_mock:
            save_mock.return_value = ("edited_v2.xlsx", "/files/1/edited_v2.xlsx")
            result = target._execute_edit_file_tool(
                1,
                {"source": "1788181323_fc48c741.xlsx", "cell_edits": [{"cell": "B2", "value": "修正値"}]},
                encrypt=False,
                loaded_files=loaded_screenshot,
                history=history,
            )
            self.assertTrue(result.get("ok"))
            self.assertEqual(result["display_name"], "1788181323_fc48c741.xlsx")
            self.assertEqual(result["url"], "/files/1/edited_v2.xlsx")

    def test_missing_source_rejected(self):
        result = target._execute_edit_file_tool(1, {"content": "x"}, encrypt=False, loaded_files=[])
        self.assertFalse(result.get("ok"))
        self.assertIn("source", result.get("error", ""))

    def test_unknown_source_rejected(self):
        result = target._execute_edit_file_tool(
            1, {"source": "nope.csv", "content": "x"}, encrypt=False,
            loaded_files=[{"send_name": "a.csv", "name": "1/x.csv", "path": "1/x.csv"}],
        )
        self.assertFalse(result.get("ok"))
        self.assertIn("見つかりません", result.get("error", ""))

    def test_missing_content_rejected(self):
        result = target._execute_edit_file_tool(
            1, {"source": "a.csv"}, encrypt=False,
            loaded_files=[{"send_name": "a.csv", "name": "1/x.csv", "path": "1/x.csv"}],
        )
        self.assertFalse(result.get("ok"))
        self.assertIn("content", result.get("error", ""))

    def test_text_file_edit_preserves_extension_and_display_name(self):
        loaded = [{"send_name": "data.csv", "name": "1/abc123.csv", "path": "1/abc123.csv"}]
        with patch("app._save_user_generated_bytes_verified") as save_mock:
            save_mock.return_value = ("data_123.csv", "/files/1/data_123.csv")
            result = target._execute_edit_file_tool(
                1, {"source": "data.csv", "content": "a,b\n1,2\n"}, encrypt=False,
                loaded_files=loaded,
            )
        self.assertTrue(result.get("ok"))
        self.assertEqual(result["display_name"], "data.csv")
        self.assertEqual(result["url"], "/files/1/data_123.csv")
        saved_bytes = save_mock.call_args[0][1]
        self.assertEqual(saved_bytes.decode("utf-8"), "a,b\n1,2\n")
        # Output filename keeps the original extension.
        made_name = save_mock.call_args[0][2]()
        self.assertTrue(made_name.endswith(".csv"))
        self.assertTrue(made_name.startswith("data_"))

    def test_xlsx_edit_applies_cell_edits_and_preserves_format(self):
        original = self._make_styled_xlsx()
        loaded = [{"send_name": "予定表.xlsx", "name": "1/abc123.xlsx", "path": "1/abc123.xlsx"}]
        with patch("app._get_file_disk_info", return_value={"exists": True, "size": len(original)}), \
             patch("app._load_user_file_bytes", return_value=original), \
             patch("app._save_user_generated_bytes_verified") as save_mock:
            save_mock.return_value = ("edited_123.xlsx", "/files/1/edited_123.xlsx")
            result = target._execute_edit_file_tool(
                1,
                {"source": "予定表.xlsx", "cell_edits": [{"cell": "B2", "value": "新値"}]},
                encrypt=False,
                loaded_files=loaded,
            )
        self.assertTrue(result.get("ok"))
        self.assertEqual(result["display_name"], "予定表.xlsx")
        saved_bytes = save_mock.call_args[0][1]
        self.assertTrue(saved_bytes.startswith(b"PK"))
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(saved_bytes))
        ws = wb["Sheet1"]
        self.assertEqual(ws["B2"].value, "新値")
        self.assertEqual(ws["C3"].value, "保持")
        self.assertEqual(ws["B2"].fill.start_color.rgb, "FFFFFF00")
        made_name = save_mock.call_args[0][2]()
        self.assertTrue(made_name.endswith(".xlsx"))

    def test_xlsx_content_mode_rejected_to_protect_format(self):
        loaded = [{"send_name": "予定表.xlsx", "name": "1/abc123.xlsx", "path": "1/abc123.xlsx"}]
        result = target._execute_edit_file_tool(
            1, {"source": "予定表.xlsx", "content": "a\tb\n1\t2\n"}, encrypt=False, loaded_files=loaded
        )
        self.assertFalse(result.get("ok"))
        self.assertIn("cell_edits", result.get("error", ""))

    def test_xlsx_cell_edits_required(self):
        loaded = [{"send_name": "予定表.xlsx", "name": "1/abc123.xlsx", "path": "1/abc123.xlsx"}]
        result = target._execute_edit_file_tool(
            1, {"source": "予定表.xlsx", "cell_edits": []}, encrypt=False, loaded_files=loaded
        )
        self.assertFalse(result.get("ok"))
        self.assertIn("cell_edits", result.get("error", ""))

    def test_apply_xlsx_cell_edits_invalid_ref_returns_error(self):
        original = self._make_styled_xlsx()
        data, errors = target._apply_xlsx_cell_edits(original, [{"cell": "INVALID", "value": "x"}])
        self.assertTrue(data)
        self.assertEqual(len(errors), 1)

    def test_binary_extension_rejected(self):
        loaded = [{"send_name": "evil.exe", "name": "1/abc123.exe", "path": "1/abc123.exe"}]
        result = target._execute_edit_file_tool(
            1, {"source": "evil.exe", "content": "x"}, encrypt=False, loaded_files=loaded
        )
        self.assertFalse(result.get("ok"))
        self.assertIn("対応していない", result.get("error", ""))

    def test_edit_file_schema_shape(self):
        schema = target._build_edit_file_tool_schema()
        self.assertEqual(schema["function"]["name"], "edit_file")
        props = schema["function"]["parameters"]["properties"]
        self.assertIn("source", props)
        self.assertIn("cell_edits", props)
        self.assertIn("content", props)
        self.assertIn("source", schema["function"]["parameters"]["required"])
        self.assertNotIn("content", schema["function"]["parameters"]["required"])

    def test_create_file_schema_mentions_edit_file(self):
        schema = target._build_create_file_tool_schema()
        self.assertIn("edit_file", schema["function"]["description"])


class XlsxCellStyleTests(unittest.TestCase):
    def _make_plain_xlsx(self):
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "売上"
        buf = BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_normalize_xlsx_color(self):
        self.assertEqual(target._normalize_xlsx_color("#FFFF00"), "FFFFFF00")
        self.assertEqual(target._normalize_xlsx_color("FFFF00"), "FFFFFF00")
        self.assertEqual(target._normalize_xlsx_color("FFFF0000"), "FFFF0000")
        self.assertEqual(target._normalize_xlsx_color("#f00"), "FFFF0000")
        self.assertIsNone(target._normalize_xlsx_color("notacolor"))
        self.assertIsNone(target._normalize_xlsx_color(""))
        self.assertIsNone(target._normalize_xlsx_color(None))
        self.assertIsNone(target._normalize_xlsx_color("#GGGGGG"))

    def test_apply_style_sets_fill_font_border_alignment_number_format(self):
        original = self._make_plain_xlsx()
        edits = [{
            "cell": "A1",
            "value": "売上(千円)",
            "style": {
                "fill": {"color": "#FFFF00"},
                "font": {"bold": True, "italic": True, "color": "#FF0000", "size": 14},
                "border": {"style": "thin", "color": "#000000"},
                "alignment": {"horizontal": "center", "vertical": "center", "wrapText": True},
                "numberFormat": "0.00%",
            },
        }]
        data, errors = target._apply_xlsx_cell_edits(original, edits)
        self.assertEqual(errors, [])
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(data))
        cell = wb["Sheet1"]["A1"]
        self.assertEqual(cell.value, "売上(千円)")
        self.assertEqual(cell.fill.start_color.rgb, "FFFFFF00")
        self.assertTrue(cell.font.bold)
        self.assertTrue(cell.font.italic)
        self.assertEqual(cell.font.color.rgb, "FFFF0000")
        self.assertEqual(cell.font.size, 14.0)
        for side in ("left", "right", "top", "bottom"):
            self.assertEqual(getattr(cell.border, side).style, "thin")
        self.assertEqual(cell.alignment.horizontal, "center")
        self.assertEqual(cell.alignment.vertical, "center")
        self.assertTrue(cell.alignment.wrap_text)
        self.assertEqual(cell.number_format, "0.00%")

    def test_style_only_edit_preserves_value(self):
        original = self._make_plain_xlsx()
        data, errors = target._apply_xlsx_cell_edits(
            original, [{"cell": "A1", "style": {"font": {"bold": True}}}]
        )
        self.assertEqual(errors, [])
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(data))
        cell = wb["Sheet1"]["A1"]
        self.assertEqual(cell.value, "売上")
        self.assertTrue(cell.font.bold)

    def test_style_preserves_unspecified_formatting(self):
        from openpyxl import Workbook
        from openpyxl.styles import PatternFill, Font
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["B2"] = "旧"
        ws["B2"].fill = PatternFill(start_color="FFFFFF00", end_color="FFFFFF00", fill_type="solid")
        ws["B2"].font = Font(bold=True, size=11)
        buf = BytesIO()
        wb.save(buf)
        original = buf.getvalue()
        # Only change the font color; fill and bold must survive.
        data, errors = target._apply_xlsx_cell_edits(
            original, [{"cell": "B2", "style": {"font": {"color": "#0000FF"}}}]
        )
        self.assertEqual(errors, [])
        from openpyxl import load_workbook
        wb2 = load_workbook(BytesIO(data))
        cell = wb2["Sheet1"]["B2"]
        self.assertEqual(cell.fill.start_color.rgb, "FFFFFF00")
        self.assertTrue(cell.font.bold)
        self.assertEqual(cell.font.size, 11.0)
        self.assertEqual(cell.font.color.rgb, "FF0000FF")

    def test_style_fill_none_clears_fill(self):
        from openpyxl import Workbook
        from openpyxl.styles import PatternFill
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "x"
        ws["A1"].fill = PatternFill(start_color="FFFFFF00", end_color="FFFFFF00", fill_type="solid")
        buf = BytesIO()
        wb.save(buf)
        original = buf.getvalue()
        data, errors = target._apply_xlsx_cell_edits(
            original, [{"cell": "A1", "style": {"fill": {"fillType": "none"}}}]
        )
        self.assertEqual(errors, [])
        from openpyxl import load_workbook
        wb2 = load_workbook(BytesIO(data))
        self.assertIsNone(wb2["Sheet1"]["A1"].fill.fill_type)

    def test_style_per_side_border_override(self):
        original = self._make_plain_xlsx()
        data, errors = target._apply_xlsx_cell_edits(
            original, [{"cell": "A1", "style": {"border": {
                "style": "thin",
                "left": {"style": "thick", "color": "00FF00"},
            }}}]
        )
        self.assertEqual(errors, [])
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(data))
        cell = wb["Sheet1"]["A1"]
        self.assertEqual(cell.border.left.style, "thick")
        self.assertEqual(cell.border.left.color.rgb, "FF00FF00")
        self.assertEqual(cell.border.right.style, "thin")
        self.assertEqual(cell.border.top.style, "thin")

    def test_style_invalid_values_report_errors_but_keep_valid_parts(self):
        original = self._make_plain_xlsx()
        data, errors = target._apply_xlsx_cell_edits(
            original, [{
                "cell": "A1",
                "style": {
                    "fill": {"color": "notacolor"},
                    "font": {"bold": True},
                },
            }]
        )
        self.assertEqual(len(errors), 1)
        self.assertIn("塗りつぶし色", errors[0])
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(data))
        cell = wb["Sheet1"]["A1"]
        self.assertTrue(cell.font.bold)

    def test_style_non_object_reports_error(self):
        original = self._make_plain_xlsx()
        data, errors = target._apply_xlsx_cell_edits(
            original, [{"cell": "A1", "style": "nope"}]
        )
        self.assertEqual(len(errors), 1)
        self.assertIn("style はオブジェクト", errors[0])

    def test_edit_file_schema_includes_cell_style(self):
        schema = target._build_edit_file_tool_schema()
        items = schema["function"]["parameters"]["properties"]["cell_edits"]["items"]
        self.assertIn("style", items["properties"])
        style_props = items["properties"]["style"]["properties"]
        for key in ("fill", "font", "border", "alignment", "numberFormat"):
            self.assertIn(key, style_props)
        # value is optional when only styling a cell.
        self.assertEqual(items["required"], ["cell"])

    def test_execute_edit_file_applies_style(self):
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "売上"
        ws["B2"] = 100
        buf = BytesIO()
        wb.save(buf)
        original = buf.getvalue()
        loaded = [{"send_name": "予定表.xlsx", "name": "1/abc123.xlsx", "path": "1/abc123.xlsx"}]
        with patch("app._get_file_disk_info", return_value={"exists": True, "size": len(original)}), \
             patch("app._load_user_file_bytes", return_value=original), \
             patch("app._save_user_generated_bytes_verified") as save_mock:
            save_mock.return_value = ("edited_123.xlsx", "/files/1/edited_123.xlsx")
            result = target._execute_edit_file_tool(
                1,
                {"source": "予定表.xlsx", "cell_edits": [
                    {"cell": "A1", "style": {"font": {"bold": True}, "fill": {"color": "#FFFF00"}}},
                    {"cell": "B2", "value": 120, "style": {"border": {"style": "thin"}}},
                ]},
                encrypt=False,
                loaded_files=loaded,
            )
        self.assertTrue(result.get("ok"))
        saved_bytes = save_mock.call_args[0][1]
        from openpyxl import load_workbook
        wb2 = load_workbook(BytesIO(saved_bytes))
        ws2 = wb2["Sheet1"]
        self.assertEqual(ws2["A1"].value, "売上")
        self.assertTrue(ws2["A1"].font.bold)
        self.assertEqual(ws2["A1"].fill.start_color.rgb, "FFFFFF00")
        self.assertEqual(ws2["B2"].value, 120)
        self.assertEqual(ws2["B2"].border.left.style, "thin")


class DocxPdfEditToolTests(unittest.TestCase):
    def _make_styled_docx(self):
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        doc.add_heading("見出し", level=1)
        p = doc.add_paragraph()
        r = p.add_run("太字テキスト")
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0, 0)
        r.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("ふつうの文章。")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_docx_numbered_extraction(self):
        data = self._make_styled_docx()
        text = target._extract_docx_as_numbered(data)
        self.assertIsNotNone(text)
        self.assertIn("[1] 見出し", text)
        self.assertIn("[2] 太字テキスト", text)
        self.assertIn("[3] ふつうの文章。", text)

    def test_docx_paragraph_edits_text_and_style(self):
        original = self._make_styled_docx()
        data, errors = target._apply_docx_paragraph_edits(original, [{
            "paragraph": 2,
            "text": "新しい太字テキスト",
            "style": {"font": {"bold": True, "color": "#0000FF"}, "alignment": "center"},
        }])
        self.assertEqual(errors, [])
        from docx import Document
        doc = Document(BytesIO(data))
        p = doc.paragraphs[1]
        self.assertEqual(p.text, "新しい太字テキスト")
        self.assertTrue(p.runs[0].bold)
        self.assertEqual(str(p.runs[0].font.color.rgb), "0000FF")

    def test_docx_style_only_edit_preserves_text(self):
        original = self._make_styled_docx()
        data, errors = target._apply_docx_paragraph_edits(original, [{
            "paragraph": 3,
            "style": {"font": {"italic": True}},
        }])
        self.assertEqual(errors, [])
        from docx import Document
        doc = Document(BytesIO(data))
        p = doc.paragraphs[2]
        self.assertEqual(p.text, "ふつうの文章。")
        self.assertTrue(p.runs[0].italic)

    def test_docx_paragraph_by_text_match(self):
        original = self._make_styled_docx()
        data, errors = target._apply_docx_paragraph_edits(original, [{
            "paragraph": "ふつうの文章",
            "text": "変更後",
        }])
        self.assertEqual(errors, [])
        from docx import Document
        doc = Document(BytesIO(data))
        self.assertEqual(doc.paragraphs[2].text, "変更後")

    def test_docx_unknown_paragraph_reports_error(self):
        original = self._make_styled_docx()
        data, errors = target._apply_docx_paragraph_edits(original, [{"paragraph": 999, "text": "x"}])
        self.assertTrue(data)
        self.assertEqual(len(errors), 1)
        self.assertIn("段落が見つかりません", errors[0])

    def test_docx_convert_to_pdf(self):
        original = self._make_styled_docx()
        pdf = target._convert_docx_to_pdf_bytes(original)
        self.assertTrue(pdf.startswith(b"%PDF"))

    def test_pdf_text_edits_replaces_and_keeps_layout(self):
        pdf = target._build_created_pdf_bytes("# テスト\n\n売上は 100 万円です。", "テスト")
        out, errors = target._apply_pdf_text_edits(pdf, [{"find": "売上は", "replace": "利益は"}])
        self.assertEqual(errors, [])
        import pymupdf
        doc = pymupdf.open(stream=out, filetype="pdf")
        self.assertIn("利益は", doc[0].get_text())

    def test_pdf_text_edits_not_found_reports_error(self):
        pdf = target._build_created_pdf_bytes("# テスト\n\n本文。", "テスト")
        out, errors = target._apply_pdf_text_edits(pdf, [{"find": "存在しない文字列", "replace": "x"}])
        self.assertTrue(out)
        self.assertEqual(len(errors), 1)
        self.assertIn("見つかりません", errors[0])

    def test_execute_docx_paragraph_edit_returns_extra_pdf(self):
        original = self._make_styled_docx()
        loaded = [{"send_name": "原稿.docx", "name": "1/abc.docx", "path": "1/abc.docx"}]
        call_count = {"n": 0}

        def fake_save(user_id, data, name_fn, encrypt):
            call_count["n"] += 1
            if call_count["n"] == 1:
                return ("edited_1.docx", "/files/1/edited_1.docx")
            return ("edited_1.pdf", "/files/1/edited_1.pdf")

        with patch("app._get_file_disk_info", return_value={"exists": True, "size": len(original)}), \
             patch("app._load_user_file_bytes", return_value=original), \
             patch("app._save_user_generated_bytes_verified", side_effect=fake_save) as save_mock:
            result = target._execute_edit_file_tool(
                1,
                {"source": "原稿.docx", "paragraph_edits": [
                    {"paragraph": 1, "text": "変更", "style": {"font": {"bold": True}}}
                ]},
                encrypt=False, loaded_files=loaded,
            )
        self.assertTrue(result.get("ok"))
        self.assertEqual(result["display_name"], "原稿.docx")
        extras = result.get("extra_files") or []
        self.assertEqual(len(extras), 1)
        self.assertTrue(extras[0]["display_name"].endswith(".pdf"))
        saved_bytes = save_mock.call_args_list[0][0][1]
        self.assertTrue(saved_bytes.startswith(b"PK"))  # docx zip container

    def test_execute_pdf_text_edit(self):
        pdf = target._build_created_pdf_bytes("# タイトル\n\n変更したい文字列です。", "テスト")
        loaded = [{"send_name": "資料.pdf", "name": "1/abc.pdf", "path": "1/abc.pdf"}]
        with patch("app._get_file_disk_info", return_value={"exists": True, "size": len(pdf)}), \
             patch("app._load_user_file_bytes", return_value=pdf), \
             patch("app._save_user_generated_bytes_verified") as save_mock:
            save_mock.return_value = ("edited_1.pdf", "/files/1/edited_1.pdf")
            result = target._execute_edit_file_tool(
                1,
                {"source": "資料.pdf", "text_edits": [{"find": "変更したい", "replace": "修正済み"}]},
                encrypt=False, loaded_files=loaded,
            )
        self.assertTrue(result.get("ok"))
        self.assertEqual(result["display_name"], "資料.pdf")
        saved_bytes = save_mock.call_args[0][1]
        self.assertTrue(saved_bytes.startswith(b"%PDF"))

    def test_edit_file_schema_includes_docx_and_pdf_edits(self):
        schema = target._build_edit_file_tool_schema()
        props = schema["function"]["parameters"]["properties"]
        self.assertIn("paragraph_edits", props)
        self.assertIn("text_edits", props)
        self.assertIn("style", props["paragraph_edits"]["items"]["properties"])
        self.assertIn("find", props["text_edits"]["items"]["properties"])
        self.assertIn("replace", props["text_edits"]["items"]["properties"])

    def test_gemini_edit_tool_accepts_paragraph_and_text_edits(self):
        app_source = read_app_source()
        idx = app_source.index("def _gemini_edit_file_tool(")
        branch = app_source[idx:idx + 4000]
        self.assertIn("cell_edits: Optional[list[dict]] = None", branch)
        self.assertIn("paragraph_edits: Optional[list[dict]] = None", branch)
        self.assertIn("text_edits: Optional[list[dict]] = None", branch)
        self.assertIn('"paragraph_edits": paragraph_edits', branch)
        self.assertIn('"text_edits": text_edits', branch)

    def test_gemini_edit_tool_declaration_has_items_for_array_properties(self):
        from google.genai import types, Client
        from typing import Optional

        def _sample_gemini_edit_file_tool(
            source: str,
            content: Optional[str] = None,
            cell_edits: Optional[list[dict]] = None,
            paragraph_edits: Optional[list[dict]] = None,
            text_edits: Optional[list[dict]] = None,
        ) -> str:
            """Sample docstring."""
            return "ok"

        _sample_gemini_edit_file_tool.__name__ = "edit_file"
        client = Client(api_key="fake")
        decl = types.FunctionDeclaration.from_callable(client=client, callable=_sample_gemini_edit_file_tool)
        for key in ("cell_edits", "paragraph_edits", "text_edits"):
            prop = decl.parameters.properties.get(key)
            self.assertIsNotNone(prop, f"Property {key} missing")
            self.assertEqual(prop.type, types.Type.ARRAY)
            self.assertIsNotNone(prop.items, f"Property {key} must have items definition")


if __name__ == "__main__":
    unittest.main()
